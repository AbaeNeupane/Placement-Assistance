# Resume parsing utility functions
import os
import re
import csv
from typing import Dict, List, Set

# Optional deps (safe import so the app won't crash if missing)
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None


# Text extraction from various file types
def _extract_text_pdf(path: str) -> str:
    if PyPDF2 is None:
        return ""
    try:
        reader = PyPDF2.PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        return ""

def _extract_text_docx(path: str) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        return ""

def _extract_text_plain(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""

def _read_all_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_text_pdf(path)
    if ext == ".docx":
        return _extract_text_docx(path)
    if ext in {".txt", ".doc"}:
        # .doc treated as plain text best-effort
        return _extract_text_plain(path)
    return ""


# Normalization and regexes
def _normalize(s: str) -> str:
    # keep tech tokens like c++, .net, email @; normalize spaces
    s = s.replace("\u00A0", " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()

EMAIL_RX = re.compile(r"[a-z0-9._%+\-]+@[a-z0-9.\-]+\.[a-z]{2,}", re.I)
PHONE_RX = re.compile(r"(\+?\d[\d \-]{7,}\d)")
YRS_RXES = [
    re.compile(r"(\d{1,2})\s*\+?\s*(?:years?|yrs?)\b", re.I),
    re.compile(r"(?:experience|exp)\s*[:\-]?\s*(\d{1,2})\s*(?:years?|yrs?)\b", re.I),
]


# Basic field extraction
def _extract_email(text: str) -> str:
    m = EMAIL_RX.search(text)
    return m.group(0) if m else ""

def _extract_phone(text: str) -> str:
    m = PHONE_RX.search(text)
    return _normalize(m.group(1)) if m else ""

def _extract_experience(text: str) -> str:
    years: List[int] = []
    for rx in YRS_RXES:
        for m in rx.finditer(text):
            try:
                years.append(int(m.group(1)))
            except Exception:
                pass
    if years:
        return str(max(years))
    m = re.search(r"(\d{1,2})\s*\+\s*(?:years?|yrs?)\s+of\s+experience", text, re.I)
    return m.group(1) if m else ""


# Name extraction
_BAD_TITLE_WORDS = {
    "resume","curriculum","vitae","cv","profile","summary",
    "developer","engineer","manager","analyst","teacher","student","intern"
}

def _clean_line_caps(s: str) -> str:
    s = _normalize(s)
    return re.sub(r"\s{2,}", " ", s)

def _looks_like_name(line: str) -> bool:
    tokens = line.strip().split()
    if not (2 <= len(tokens) <= 4):
        return False
    # Every token: starts uppercase, letters/dots/hyphens allowed
    for t in tokens:
        if not re.match(r"^[A-Z][A-Za-z\.'\-]*$", t):
            return False
    if {t.lower() for t in tokens} & _BAD_TITLE_WORDS:
        return False
    return True

def _extract_name_from_docx(path: str) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(path)
    except Exception:
        return ""

    # 1. Paragraphs: prioritize Title/Heading styles in top 20
    heading_hits, plain_hits = [], []
    for p in d.paragraphs[:20]:
        text = _clean_line_caps(p.text)
        if not text:
            continue
        style = (p.style.name if p.style else "") or ""
        # Labeled "Name: ..."
        m = re.search(r'\b(full\s+name|name)\s*[:\-–]\s*([A-Za-z][A-Za-z\.\'\-\s]{2,60})', text, re.I)
        if m:
            cand = _clean_line_caps(m.group(2))
            if _looks_like_name(cand):
                return cand
        if _looks_like_name(text):
            if any(s in style.lower() for s in ("title", "heading")):
                heading_hits.append(text)
            else:
                plain_hits.append(text)
    if heading_hits:
        return heading_hits[0]
    if plain_hits:
        return plain_hits[0]

    # 2. Tables: designer headers (top 2 tables, first 2 rows)
    for tbl in d.tables[:2]:
        for row in tbl.rows[:2]:
            for cell in row.cells:
                t = _clean_line_caps(cell.text)
                if not t:
                    continue
                m = re.search(r'\b(full\s+name|name)\s*[:\-–]\s*([A-Za-z][A-Za-z\.\'\-\s]{2,60})', t, re.I)
                if m:
                    cand = _clean_line_caps(m.group(2))
                    if _looks_like_name(cand):
                        return cand
                if _looks_like_name(t):
                    return t
    return ""

def _guess_name_from_email(email: str) -> str:
    if not email:
        return ""
    local = email.split("@", 1)[0]
    parts = re.split(r"[._\-]+", local)
    parts = [p for p in parts if p.isalpha()]
    if 1 <= len(parts) <= 3:
        return " ".join(w.capitalize() for w in parts)
    return ""

def _guess_name(text: str, email: str) -> str:
  # Scan top 30 lines for a name-like line
    for raw in text.splitlines()[:30]:
        line = _clean_line_caps(raw)
        if len(line) < 3:
            continue
        if re.search(r"(curriculum vitae|resume|cv)", line, re.I):
            continue
        if re.search(r"(teacher|developer|engineer|manager|analyst|student|intern)", line, re.I):
            continue
        if _looks_like_name(line):
            return line
    return _guess_name_from_email(email)


# Designation extraction job titles
TITLE_KEYWORDS = {
    # teaching
    "teacher","primary teacher","secondary teacher","class teacher",
    "english teacher","math teacher","science teacher",
    # tech / business
    "software developer","software engineer","backend developer",
    "fullstack developer","data scientist","data engineer",
    "machine learning engineer","devops engineer","system administrator",
    "it support","business analyst","business intelligence analyst",
    "product manager","qa engineer","quality assurance engineer",
    "frontend developer","data analyst","project manager",
    "sales executive","store manager","retail associate",
    # generic title words
    "developer","engineer","manager","analyst","student","intern",
}

SEP_RX = re.compile(r"\s*[|\-–—]\s*")  # splits "Name | Title" or "Name - Title"

def _titlecase_phrase(s: str) -> str:
    parts = [p.strip() for p in re.split(r"\s+", s.strip()) if p.strip()]
    out = []
    for p in parts:
        if len(p) <= 3 and p.isupper():
            out.append(p)
        else:
            out.append(p.capitalize())
    return " ".join(out)

# Extract designation around name in text lines
def _extract_designation_from_text_around_name(lines: List[str], name: str) -> str:
#   Look around the name line for title-like chunks (handles separators).
    name_l = (name or "").lower()
    for idx, raw in enumerate(lines[:20]):
        line = _normalize(raw)
        if not line:
            continue
        if name_l and name_l in line.lower() and re.search(SEP_RX, line):
            segs = [s.strip() for s in re.split(SEP_RX, line) if s.strip()]
            try:
                i = next(i for i, s in enumerate(segs) if name_l in s.lower())
            except StopIteration:
                i = None
            candidates = []
            if i is not None:
                if i + 1 < len(segs):
                    candidates.append(segs[i + 1])
                if i - 1 >= 0:
                    candidates.append(segs[i - 1])
            else:
                candidates = segs

            for cand in candidates:
                lc = cand.lower()
                if any(k in lc for k in TITLE_KEYWORDS) or 1 <= len(cand.split()) <= 5:
                    return _titlecase_phrase(cand)

        # If the next line looks like a title
        if name_l and name_l in line.lower() and idx + 1 < len(lines):
            nxt = _normalize(lines[idx + 1])
            if nxt and (any(k in nxt.lower() for k in TITLE_KEYWORDS) or 1 <= len(nxt.split()) <= 6):
                return _titlecase_phrase(nxt)

    # Fallback: scan top lines for something that *looks* like a title
    for raw in lines[:15]:
        t = _normalize(raw)
        if not t:
            continue
        if any(k in t.lower() for k in TITLE_KEYWORDS):
            return _titlecase_phrase(t)
    return ""
# Designation extraction from DOCX
def _extract_designation_from_docx(path: str, name: str) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(path)
    except Exception:
        return ""
    # 1) Paragraphs near the top
    top_paras = [p.text for p in d.paragraphs[:25]]
    title = _extract_designation_from_text_around_name(top_paras, name)
    if title:
        return title
    # 2) Top table cells (designer headers)
    for tbl in d.tables[:2]:
        cells = []
        for row in tbl.rows[:2]:
            for cell in row.cells:
                cells.append(cell.text)
        if cells:
            title = _extract_designation_from_text_around_name(cells, name)
            if title:
                return title
    return ""

# Skills extraction
KEY_SKILLS_HDR_RX = re.compile(r"\bkey\s*skills?\s*:\s*", re.I)
PREF_DESIG_HDR_RX = re.compile(r"\b(my\s+preferred\s+designation)\s*:\s*", re.I)

def _split_tokens(s: str) -> List[str]:
    """Split a skills line by | , ; / while keeping tech tokens."""
    s = s.replace("|", ",")
    parts = re.split(r"[,\;/]+", s)
    return [p.strip() for p in parts if p.strip()]

def _extract_preferred_designation(text: str) -> str:
    """
    Find 'My preferred Designation:' line and return ALL listed roles,
    e.g., 'Software Developer, Backend Developer, ...'
    """
    for line in text.splitlines():
        m = PREF_DESIG_HDR_RX.search(line)
        if m:
            tail = line[m.end():].strip()
            # split by commas and 'and'
            raw_roles = re.split(r",|\band\b", tail, flags=re.I)
            roles = [r.strip(" .") for r in raw_roles if r and len(r.strip()) >= 2]
            if roles:
                # Title-case each word (preserve common acronyms)
                def _tc_word(w: str) -> str:
                    return w if (len(w) <= 3 and w.isupper()) else w.title()
                def _tc_phrase(ph: str) -> str:
                    return " ".join(_tc_word(x) for x in ph.split())
                # JOIN ALL roles (your request)
                return ", ".join(_tc_phrase(r) for r in roles)
    return ""

def _extract_key_skills_block(text: str) -> List[str]:
    """
    Pull skills specifically from the 'Key skills:' section, including
    subsection lines like 'Frontend: Html|Css|Js|Tailwind|React' and
    'Backend: NodeJs|Express|Python|Flask|Django'.
    Stops when an empty line or a new labeled section is reached.
    """
    lines = text.splitlines()
    skills: List[str] = []
    capturing = False
    for raw in lines:
        line = raw.strip()
        if not capturing:
            if KEY_SKILLS_HDR_RX.search(line):
                tail = KEY_SKILLS_HDR_RX.sub("", line).strip()
                if tail:
                    tail = re.sub(r"^\s*(frontend|backend|database)\s*:\s*", "", tail, flags=re.I)
                    skills.extend(_split_tokens(tail))
                capturing = True
            continue
        if not line or PREF_DESIG_HDR_RX.search(line) or re.match(r"^\s*(experience|projects?|education)\s*:", line, re.I):
            break
        line = re.sub(r"^\s*(frontend|backend|database)\s*:\s*", "", line, flags=re.I)
        skills.extend(_split_tokens(line))

    # normalize unique (case-insensitive)
    seen, out = set(), []
    for s in skills:
        t = s.strip()
        if t and t.lower() not in seen:
            out.append(t)
            seen.add(t.lower())
    return out


# Build skills vocabulary from jobs_info.csv
def _build_skills_vocab_from_jobs_csv() -> Set[str]:
    vocab: Set[str] = set()
    base_dir = os.path.dirname(os.path.abspath(__file__))
    csv_path = os.path.join(base_dir, "jobs_info.csv")
    if os.path.exists(csv_path):
        try:
            with open(csv_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.DictReader(f)
                for r in reader:
                    raw = (r.get("Key Skills") or "").lower()
                    for token in re.split(r"[,\;/\|]+", raw):
                        token = token.strip()
                        if token:
                            vocab.add(token)
        except Exception:
            pass

    if not vocab:
        vocab.update({
            # teaching
            "lesson planning", "classroom management", "curriculum design",
            "english", "mathematics", "science",
            # tech
            "python", "flask", "mongodb", "rest api", "git",
            "html", "css", "javascript", "react", "tailwind", "nodejs", "express", "django", "mysql",
            # retail
            "customer service", "sales", "cash handling", "inventory",
        })
    return vocab

def _extract_skills(text: str, vocab: Set[str]) -> List[str]:
    norm = " " + _normalize(text).lower() + " "
    found: List[str] = []
    for phrase in sorted(vocab, key=lambda x: (-len(x), x)):
        if (" " + phrase + " ") in norm:
            found.append(phrase)
    # de-dup preserve order
    seen, uniq = set(), []
    for s in found:
        if s not in seen:
            uniq.append(s); seen.add(s)
    return uniq[:50]


# API-compatible resume parsing
def parse_resume(path: str) -> Dict[str, object]:
    """
    Returns a dict compatible with /api/parse_cv in app.py:
      name, email, phone, designation, experience, skills(list), education
    """
    ext = os.path.splitext(path)[1].lower()
    text = _read_all_text(path)
    norm_text = _normalize(text)

    # Basic fields
    email = _extract_email(norm_text)
    phone = _extract_phone(norm_text)

    # Name: DOCX-special handling, then fallbacks
    name = ""
    if ext == ".docx":
        name = _extract_name_from_docx(path)
    if not name:
        name = _guess_name(norm_text, email)

# Designation: prefer 'My preferred Designation:' line
    designation = _extract_preferred_designation(text)

    # Fallbacks
    if not designation and ext == ".docx":
        designation = _extract_designation_from_docx(path, name)
    if not designation:
        # generic text fallback (labels/headlines) - use the text-around-name helper
        designation = _extract_designation_from_text_around_name(text.splitlines(), name)

    # Experience
    experience = _extract_experience(norm_text)

    # Skills: prefer 'Key skills:' block, then vocab-based generic skills
    block_skills = _extract_key_skills_block(text)
    vocab = _build_skills_vocab_from_jobs_csv()
    generic_skills = _extract_skills(norm_text, vocab)

    # Merge with preference for block skills first
    merged: List[str] = []
    seen = set()
    for s in block_skills + generic_skills:
        sl = s.lower()
        if sl not in seen:
            merged.append(s)
            seen.add(sl)

    # De-contaminate: remove designation terms / generic title words
    desig_l = (designation or "").lower()
    desig_tokens = {t for t in re.split(r"[ /,;|\-–—]+", desig_l) if t}
    TITLE_WORDS = {"developer", "engineer", "manager", "analyst", "teacher", "student", "intern"}
    cleaned = []
    for s in merged:
        sl = s.lower()
        if sl in desig_tokens:
            continue
        if sl in TITLE_WORDS:
            continue
        if desig_l and (sl in desig_l or desig_l in sl):
            continue
        cleaned.append(s)

    # Education snippet (best-effort)
    edu_match = re.search(r"(education|academic\s+qualifications|qualifications)[:\n\r ]+(.{0,300})",
                          text, re.I | re.S)
    education = ""
    if edu_match:
        snippet = edu_match.group(2)
        education = _normalize(re.sub(r"\s+", " ", snippet))[:150]

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "designation": designation,          
        "experience": experience,            
        "skills": cleaned,                  
        "education": education,
    }
